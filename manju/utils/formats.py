"""manju format I/O — Multi-format reader/writer for Excel, Word, PDF, JSON, Markdown."""

import json, os
from html import escape

from manju.pipeline.storyboard_schema import (
    duration_label,
    get_audio,
    get_characters,
    get_prompt,
    get_scene_heading,
    get_spoken_text,
    get_visual,
)

# ── Excel (openpyxl) ─────────────────────────────────────────────────────────

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False


def write_xlsx(data: dict, path: str, sheet_name: str = "Sheet1"):
    if not HAS_EXCEL:
        raise ImportError("openpyxl not installed")

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name

    hdr_font = Font(name="宋体", size=10, bold=False, color="000000")
    hdr_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    cell_font = Font(name="宋体", size=10, bold=False, color="000000")
    cell_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    cell_align = Alignment(vertical="top", wrap_text=True)
    gray_side = Side(style="thin", color="B7B7B7")
    border = Border(left=gray_side, right=gray_side, top=gray_side, bottom=gray_side)

    rows = _flatten_for_excel(data)
    if rows:
        headers = list(rows[0].keys())
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=1, column=col, value=h)
            c.font, c.fill, c.alignment, c.border = hdr_font, hdr_fill, hdr_align, border
        for ri, rd in enumerate(rows, 2):
            for col, key in enumerate(headers, 1):
                v = rd.get(key, "")
                if isinstance(v, (list, dict)):
                    v = json.dumps(v, ensure_ascii=False)
                c = ws.cell(row=ri, column=col, value=v)
                c.font, c.fill, c.alignment, c.border = cell_font, cell_fill, cell_align, border
        for col in range(1, len(headers) + 1):
            mx = len(str(headers[col - 1]))
            for row in range(2, len(rows) + 2):
                val = str(ws.cell(row=row, column=col).value or "")
                cn = sum(1 for c in val if '\u4e00' <= c <= '\u9fff')
                mx = max(mx, cn * 2 + (len(val) - cn))
            ws.column_dimensions[get_column_letter(col)].width = min(mx + 4, 60)
    wb.save(path)


def _flatten_for_excel(data: dict) -> list[dict]:
    rows = []
    if "scenes" in data:
        for scene in data.get("scenes", []):
            for shot in scene.get("shots", []):
                rows.append({
                    "场景": scene.get("scene_id", ""),
                    "场景标题": get_scene_heading(scene),
                    "场景母版": scene.get("scene_template", ""),
                    "氛围": scene.get("visual_mood", ""),
                    "镜号": shot.get("shot_id", ""),
                    "景别": get_visual(shot, "shot_type"),
                    "构图": get_visual(shot, "composition"),
                    "构图情感": get_visual(shot, "composition_emotion"),
                    "运镜": get_visual(shot, "camera_movement"),
                    "时长": duration_label(shot),
                    "画面描述": get_visual(shot, "description"),
                    "对白": get_spoken_text(shot),
                    "音效": get_audio(shot, "sound_music"),
                    "色调": get_visual(shot, "color_tone"),
                    "中文生图提示词": get_prompt(shot, "image_cn"),
                    "英文生图提示词": get_prompt(shot, "image_en"),
                    "视频提示词(中文)": get_prompt(shot, "video_cn"),
                    "视频提示词(英文)": get_prompt(shot, "video_en"),
                })
    elif "lines" in data:
        for line in data.get("lines", []):
            rows.append({
                "镜头": line.get("shot_id", ""), "角色": line.get("character", ""),
                "台词": line.get("text", ""), "情绪": line.get("emotion", ""),
                "语速": line.get("speed", ""), "声调": line.get("pitch", ""),
                "音量": line.get("volume", ""), "音色描述": line.get("voice_description", ""),
                "Edge音色": line.get("voice_edge", ""), "API音色": line.get("voice_api", ""),
            })
    elif "shots" in data:
        for shot in data.get("shots", []):
            rows.append({
                "镜号": shot.get("shot_id", ""), "时长": shot.get("duration", ""),
                "运镜类型": shot.get("camera_movement_en", shot.get("camera_movement_original", "")),
                "视频提示词": shot.get("video_prompt", ""),
                "画面描述": shot.get("visual_description", ""),
            })
    elif "characters" in data and not rows:
        for char in data.get("characters", []):
            rows.append({"角色名": char.get("name", ""), "定位": char.get("role", ""),
                         "视觉锚定": char.get("visual_anchor", char.get("anchor_description", ""))})
    if not rows:
        rows.append({"内容": json.dumps(data, ensure_ascii=False, indent=2)})
    return rows


# ── Word (python-docx) ────────────────────────────────────────────────────────

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── Shared docx style helper ─────────────────────────────────────────────────

BODY = Pt(10) if HAS_DOCX else None
SMALL = Pt(9) if HAS_DOCX else None

def _shade_cell(cell, color: str):
    """Apply background shading to a cell."""
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(shd)

def _set_cell(cell, text: str, bold: bool = False, size=SMALL,
              shade: str = ""):
    """Set cell text with SimSun font, no bold, controlled formatting."""
    from docx.oxml.ns import qn
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(str(text))
    run.font.size = size
    run.bold = False  # never bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)  # always dark text
    if shade:
        _shade_cell(cell, shade)

def _make_table(doc, headers: list[str], rows: list[list[str]]):
    """Create a formatted table with borders, shading, SimSun font."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    n = min(len(rows) + 1, 100)
    table = doc.add_table(rows=n, cols=len(headers))
    table.autofit = True
    # Set table borders (thin, single line)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '999999')
        borders.append(el)
    tblPr.append(borders)
    # Header row — light blue bg, dark text, SimSun
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=False, size=SMALL, shade="D6E4F0")
    # Data rows — white bg, SimSun, no bold
    for ri, row_data in enumerate(rows[:n-1]):
        for ci, val in enumerate(row_data):
            _set_cell(table.rows[ri+1].cells[ci], val, bold=False, size=SMALL)
    return table


def write_docx(data: dict, path: str, title: str = ""):
    if not HAS_DOCX:
        raise ImportError("python-docx not installed")

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = BODY
    style.paragraph_format.space_after = Pt(4)
    # Set east-asian font too
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if not title and "title" in data:
        title = data["title"]
    h = doc.add_heading(title or "manju 输出", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Characters (v2 creative_bible or legacy top-level list)
    characters = get_characters(data)
    if characters:
        doc.add_heading("角色列表", level=2)
        _make_table(doc, ["角色", "定位", "视觉锚定"],
                    [[c.get("name", ""), c.get("role", ""),
                      c.get("visual_anchor", c.get("anchor_description", ""))]
                     for c in characters])

    # Storyboard scenes
    if "scenes" in data:
        doc.add_heading("分镜脚本", level=2)
        for scene in data["scenes"]:
            sid = scene.get("scene_id", "?")
            doc.add_heading(f"场景 {sid}：{get_scene_heading(scene)}", level=3)
            if scene.get("visual_mood"):
                p = doc.add_paragraph()
                p.add_run("氛围：").bold = True
                p.add_run(scene["visual_mood"])
            shots = scene.get("shots", [])
            if shots:
                _make_table(doc, ["镜号","景别","构图","运镜","时长","画面描述","对白"],
                            [[
                                str(s.get("shot_id", "")),
                                get_visual(s, "shot_type"),
                                get_visual(s, "composition"),
                                get_visual(s, "camera_movement"),
                                duration_label(s),
                                get_visual(s, "description"),
                                get_spoken_text(s),
                            ] for s in shots])
            doc.add_paragraph()

    # Voice
    if "lines" in data:
        doc.add_heading("配音脚本", level=2)
        lines = data["lines"]
        if lines:
            _make_table(doc, ["镜头","角色","台词","情绪","音色描述","Edge音色","API音色","语速","声调","音量"],
                       [[str(l.get(k,"")) for k in ["shot_id","character","text","emotion","voice_description","voice_edge","voice_api","speed","pitch","volume"]]
                        for l in lines])

    # Video
    if "shots" in data and data.get("shots") and any(
        k in data["shots"][0] for k in ("video_prompt_cn", "video_prompt_en")):
        doc.add_heading("视频提示词", level=2)
        def _prompt_cell(s):
            cn = s.get("video_prompt_cn", "")
            en = s.get("video_prompt_en", "")
            return f"{cn}\n\n{en}" if cn and en else cn or en
        _make_table(doc, ["镜号","景别","时长","运镜","对白","中文提示词 / English Prompt"],
                   [[str(s.get("shot_id","")), str(s.get("shot_type","")), str(s.get("duration","")),
                     f'{s.get("camera_movement_original","")} → {s.get("camera_movement_en","")}',
                     str(s.get("dialogue","")), _prompt_cell(s)]
                    for s in data["shots"]])

    doc.save(path)


# ── PDF (HTML → weasyprint) ──────────────────────────────────────────────────

def write_pdf(data: dict, path: str, title: str = ""):
    if not title and "title" in data:
        title = data["title"]

    reportlab_error = None
    try:
        from manju.utils.reportlab_pdf import write_data_pdf

        write_data_pdf(data, path, title)
        return
    except Exception as exc:
        reportlab_error = exc

    html = _build_pdf_html(data, title)
    try:
        from weasyprint import HTML

        def blocked_fetcher(url, *args, **kwargs):
            raise ValueError(f"PDF external resource blocked: {url}")

        HTML(string=html, url_fetcher=blocked_fetcher).write_pdf(path)
    except Exception as weasy_error:
        raise RuntimeError(
            "PDF export failed with both ReportLab and WeasyPrint: "
            f"ReportLab={reportlab_error}; WeasyPrint={weasy_error}"
        ) from weasy_error


def _build_pdf_html(data: dict, title: str) -> str:
    css = """
    @page { margin: 1.5cm; size: A4; }
    body { font-family: 'Microsoft YaHei','SimHei',sans-serif; font-size: 10pt; color: #222; }
    h1 { text-align: center; font-size: 16pt; margin-bottom: 4px; }
    .subtitle { text-align: center; font-size: 9pt; color: #666; margin-bottom: 12px; }
    table { border-collapse: collapse; width: 100%; margin: 8px 0; font-size: 9pt; }
    th { background: #2F5496; color: white; padding: 4px 6px; text-align: left; }
    td { border: 1px solid #ddd; padding: 3px 6px; vertical-align: top; }
    """
    if "lines" in data:
        return _html_voice(data, title, css)
    elif "shots" in data and data["shots"] and "video_prompt" in str(data.get("shots", [{}])[0]):
        return _html_video(data, title, css)
    return f"<html><head><meta charset=\"utf-8\"><style>{css}</style></head><body><h1>{escape(str(title))}</h1></body></html>"


def _html_voice(data, title, css):
    lines = data.get("lines", [])
    sl = {0.5:"极慢",0.7:"缓慢",0.8:"较慢",1.0:"正常",1.2:"稍快",1.4:"较快",1.5:"快",1.6:"很快",1.7:"极快"}
    pl = {2:"极低沉",3:"低沉",4:"偏低",5:"中等",7:"偏高",8:"高",9:"很高",10:"极高"}
    rows = "".join(
        f'<tr><td>{escape(str(l.get("shot_id","")))}</td><td>{escape(str(l.get("character","")))}</td>'
        f'<td>{escape(str(l.get("text","")))}</td><td>{escape(str(l.get("emotion","")))}</td>'
        f'<td>{escape(str(l.get("speed",1.0)))}（{escape(str(sl.get(l.get("speed",1.0),"")))}）</td>'
        f'<td>{escape(str(l.get("pitch",5)))}（{escape(str(pl.get(l.get("pitch",5),"")))}）</td>'
        f'<td>{escape(str(l.get("volume","")))}/10</td></tr>'
        for l in lines)
    return f'<html><head><meta charset="utf-8"><style>{css}</style></head><body><h1>配音脚本 — {escape(str(title))}</h1><p class="subtitle">共 {len(lines)} 句对白</p><table><tr><th>镜头</th><th>角色</th><th>台词</th><th>情绪</th><th>语速</th><th>声调</th><th>音量</th></tr>{rows}</table></body></html>'


def _html_video(data, title, css):
    shots = data.get("shots", [])
    rows = "".join(
        f'<tr><td>{escape(str(s.get("shot_id","")))}</td><td>{escape(str(s.get("shot_type","")))}</td>'
        f'<td>{escape(str(s.get("duration","")))}</td>'
        f'<td>{escape(str(s.get("camera_movement_original","")))} → {escape(str(s.get("camera_movement_en","")))}</td>'
        f'<td>{escape(str(s.get("dialogue","")))}</td>'
        f'<td>{escape(str(s.get("video_prompt_cn", "")))}<br><br>{escape(str(s.get("video_prompt_en", "")))}</td></tr>'
        for s in shots)
    return f'<html><head><meta charset="utf-8"><style>{css}</style></head><body><h1>视频提示词 — {escape(str(title))}</h1><p class="subtitle">共 {len(shots)} 个镜头</p><table><tr><th>镜头</th><th>景别</th><th>时长</th><th>运镜</th><th>对白</th><th>中文提示词 / English Prompt</th></tr>{rows}</table></body></html>'


# ── Input readers ─────────────────────────────────────────────────────────────

def read_input(path: str) -> dict | str | None:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".json":
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    elif ext in (".txt", ".md"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".docx":
        if not HAS_DOCX:
            raise ImportError("python-docx needed")
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == ".xlsx":
        if not HAS_EXCEL:
            raise ImportError("openpyxl needed")
        wb = __import__('openpyxl').load_workbook(path)
        text = ""
        for sn in wb.sheetnames:
            ws = wb[sn]
            text += f"\n=== {sn} ===\n"
            for row in ws.iter_rows(values_only=True):
                text += " | ".join(str(c) if c else "" for c in row) + "\n"
        return text
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        return None
