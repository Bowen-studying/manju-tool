# ── Use guide ─────────────────────────────────────────────────────────────────

import os
import sys
from html import escape

def write_use_guide(output_dir: str, files: dict, title: str = ""):
    """Generate a user guide PDF+DOCX — no product names, no ads."""

    css = """
    @page { margin: 1.5cm; size: A4; }
    body { font-family: 'SimSun','SimHei','Microsoft YaHei',sans-serif; font-size: 10pt; color: #222; line-height: 1.7; }
    h1 { font-size: 16pt; margin-bottom: 4px; text-align: center; }
    h2 { font-size: 13pt; margin-top: 18px; margin-bottom: 6px; border-bottom: 1px solid #bbb; padding-bottom: 4px; }
    h3 { font-size: 11pt; margin-top: 12px; margin-bottom: 4px; }
    table { border-collapse: collapse; width: 100%; margin: 8px 0 12px; font-size: 9pt; }
    th { background: #D6E4F0; padding: 5px 8px; text-align: left; }
    td { border: 1px solid #ccc; padding: 4px 8px; vertical-align: top; }
    ul, ol { margin: 4px 0 4px 20px; }
    li { margin: 2px 0; }
    """

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{css}</style></head><body>
<h1>使用指南</h1>
<p style="text-align:center;color:#666;font-size:9pt">以下为生成的全部交付物，请按顺序使用。</p>

<h2>文件清单</h2>
<table>
<tr><th>文件</th><th>用途</th><th>下一步</th></tr>
"""
    if "storyboard_xlsx" in files:
        html += f'<tr><td>{escape(str(files["storyboard_xlsx"]))}</td><td>分镜表</td><td>打开查看和编辑，每镜含景别、构图、运镜、画面描述、中英文生图和视频提示词</td></tr>'
    if "voice_pdf" in files:
        html += f'<tr><td>{escape(str(files["voice_pdf"]))}</td><td>配音脚本</td><td>按情绪、语速、声调、音量参数进行配音</td></tr>'
    if "video_pdf" in files:
        html += f'<tr><td>{escape(str(files["video_pdf"]))}</td><td>视频提示词</td><td>逐镜使用中英文提示词生成视频片段</td></tr>'
    html += '</table>'

    html += '''
<h2>制作流程</h2>

<h3>第一步：确认分镜</h3>
<p>打开分镜表格，逐镜检查画面描述、对白、景别是否符合预期。需要修改的直接在表格中编辑。</p>

<h3>第二步：生成角色图和场景图</h3>
<p>使用分镜表中的生图提示词，逐镜生成静态图。第1张确立风格后，后续镜头复用统一的视觉锚定描述以确保角色一致。</p>

<h3>第三步：图生视频</h3>
<p>打开视频提示词，每个镜头标注了：</p>
<ul>
<li>运镜方式 — 已自动映射对应的参数</li>
<li>建议时长</li>
<li>完整的中英文视频提示词</li>
</ul>
<p>将生成的静态图配合视频提示词，逐镜生成动态视频片段。</p>

<h3>第四步：配音</h3>
<p>打开配音脚本，每句对白包含：</p>
<ul>
<li>情绪标签</li>
<li>语速、声调、音量参数</li>
<li>音色描述</li>
</ul>
<p>按照脚本中的参数逐句生成配音文件。</p>

<h3>第五步：剪辑合成</h3>
<p>将生成的视频片段和配音导入剪辑软件：</p>
<ol>
<li>按镜头顺序排列视频片段</li>
<li>添加配音并对齐节奏</li>
<li>添加音效（参考分镜表的音效列）</li>
<li>添加字幕（黑体白色描边，底部居中）</li>
<li>统一调色</li>
<li>导出成品</li>
</ol>
</body></html>'''

    generated = {}
    pdf_path = os.path.join(output_dir, "使用指南.pdf")
    try:
        from manju.utils.reportlab_pdf import write_guide_pdf

        write_guide_pdf(pdf_path, files, title)
        print(f"   📕 使用指南.pdf → {pdf_path}")
        generated["pdf"] = pdf_path
    except Exception as e:
        print(f"   ⚠ ReportLab 不可用，改用 WeasyPrint: {e}", file=sys.stderr)

    if "pdf" not in generated:
        try:
            from weasyprint import HTML

            def blocked_fetcher(url, *args, **kwargs):
                raise ValueError(f"external resource blocked: {url}")
            HTML(string=html, url_fetcher=blocked_fetcher).write_pdf(pdf_path)
            print(f"   📕 使用指南.pdf → {pdf_path}")
            generated["pdf"] = pdf_path
        except Exception as e:
            print(f"   ⚠ 使用指南 PDF: {e}", file=sys.stderr)

    try:
        from docx import Document
        doc = Document()
        doc.add_heading("使用指南", 0)
        doc.add_paragraph("以下为生成的全部交付物，请按顺序使用。")
        doc.add_heading("文件清单", level=1)
        table = doc.add_table(rows=1, cols=3)
        for cell, value in zip(table.rows[0].cells, ("文件", "用途", "下一步")):
            cell.text = value
        descriptions = {
            "storyboard_xlsx": ("分镜表", "检查画面、对白、提示词"),
            "voice_pdf": ("配音脚本", "按角色和情绪生成配音"),
            "video_pdf": ("视频提示词", "逐镜生成视频素材"),
        }
        for key, path in files.items():
            if key not in descriptions:
                continue
            row = table.add_row().cells
            row[0].text = str(path)
            row[1].text, row[2].text = descriptions[key]
        for heading, body in [
            ("确认分镜", "检查镜头顺序、人物一致性、对白、景别与时长。"),
            ("生成图片", "按镜头生成静态图片；内容变化后缓存会自动失效。"),
            ("生成视频", "使用参考图和视频提示词逐镜生成视频素材。"),
            ("配音", "按角色固定音色、情绪、语速和音量生成音频。"),
            ("剪辑合成", "在剪辑软件中对齐视频、配音、字幕和音效后导出成片。"),
        ]:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(body)
        docx_path = os.path.join(output_dir, "使用指南.docx")
        doc.save(docx_path)
        generated["docx"] = docx_path
        print(f"   📘 使用指南.docx → {docx_path}")
    except Exception as e:
        print(f"   ⚠ 使用指南 Word: {e}", file=sys.stderr)
    return generated
